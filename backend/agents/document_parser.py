"""
文档解析智能体 - 负责解析上传的招标文件并生成结构化的Markdown文档
"""
from typing import Dict, Any, List
import os
import json
import binascii
import logging
from pathlib import Path
from .base import BaseAgent, AgentContext, AgentResponse


class DocumentParserAgent(BaseAgent):
    def __init__(self):
        super().__init__("document_parser")
        self.logger = logging.getLogger(__name__)
    
    def get_system_prompt(self) -> str:
        return """您是一个专业的文档解析智能体，专门负责解析招标文件（PDF/DOC/DOCX）并识别文档结构。

您的主要职责：
1. 读取上传的招标文件
2. 识别文档结构（章节标题、内容、页码等）
3. 提取文档元数据
4. 将解析结果保存到wiki文件夹中

解析要求：
- 准确识别章节标题和层级关系
- 提取文档中的表格、列表等结构化信息
- 保持原文档的逻辑结构
- 生成结构化的JSON格式输出

输出格式要求：
{
  "filename": "原文件名",
  "document_type": "招标文件类型",
  "structure": [
    {
      "level": 1,
      "title": "章节标题",
      "content": "章节内容",
      "page": "页码",
      "subsections": []
    }
  ],
  "metadata": {
    "total_pages": "总页数",
    "creation_date": "创建日期",
    "file_size": "文件大小"
  }
}"""
    
    async def execute(self, context: AgentContext) -> AgentResponse:
        try:
            parsed_documents = []
            files_to_create: List[Dict[str, Any]] = []
            aggregated_extracted: Dict[str, Any] = {}
            
            for file_info in context.uploaded_files:
                filename = file_info.get('name', 'unknown')
                file_content = file_info.get('content', '')
                file_type = file_info.get('type', '')
                self.logger.info(f"[Upload] Received file: name={filename}, type={file_type}, content_len={len(file_content)}")
                
                if not file_content:
                    continue
                
                # 保存上传的文件到uploads文件夹
                uploads_dir = "/root/project/git/project-agent/uploads"
                os.makedirs(uploads_dir, exist_ok=True)
                # 使用安全的文件名（移除特殊字符）
                safe_filename = ''.join(c for c in filename if c.isalnum() or c in (' ', '.', '-', '_')).rstrip()
                upload_path = os.path.join(uploads_dir, safe_filename)
                
                # 处理文件内容 - 支持base64编码和纯文本
                import base64
                try:
                    # 首先尝试解码base64
                    try:
                        # 检查是否为shell命令而非真实base64内容
                        if file_content.startswith('$(') and file_content.endswith(')'):
                            print(f"Warning: File content appears to be a shell command, not base64: {file_content[:50]}...")
                            raise ValueError("File content is a shell command, not actual file content")
                        
                        # 确保base64字符串正确填充
                        padding = len(file_content) % 4
                        if padding:
                            file_content += '=' * (4 - padding)
                        
                        decoded_content = base64.b64decode(file_content)
                        
                        # Determine if this is a binary or text file
                        file_ext = safe_filename.lower().split('.')[-1] if '.' in safe_filename else ''
                        is_binary = file_ext in ['pdf', 'docx', 'doc', 'xlsx', 'pptx']
                        
                        if is_binary:
                            with open(upload_path, 'wb') as f:
                                f.write(decoded_content)
                            self.logger.info(f"[Upload] Saved binary file -> {upload_path} ({len(decoded_content)} bytes)")
                        else:
                            with open(upload_path, 'w', encoding='utf-8') as f:
                                f.write(decoded_content.decode('utf-8'))
                            self.logger.info(f"[Upload] Saved text file -> {upload_path} (chars={len(decoded_content.decode('utf-8'))})")
                    except (binascii.Error, UnicodeDecodeError, ValueError) as e:
                        # 如果不是base64，直接保存为纯文本（但记录错误）
                        self.logger.warning(f"[Upload] Base64 decode failed for {filename}: {e}")
                        print(f"Content length: {len(file_content)}")
                        print(f"Content sample: {file_content[:100]}...")
                        
                        # 对于shell命令，跳过处理
                        if file_content.startswith('$(') and file_content.endswith(')'):
                            print(f"Skipping file {filename} as it contains shell command instead of actual content")
                            continue
                            
                        with open(upload_path, 'w', encoding='utf-8') as f:
                            f.write(file_content)
                        self.logger.info(f"[Upload] Saved raw text file -> {upload_path} (chars={len(file_content)})")
                        
                except Exception as e:
                    self.logger.error(f"[Upload] Error saving file {filename}: {e}")
                    print(f"Content length: {len(file_content)}")
                    print(f"Content sample: {file_content[:100]}...")
                    continue
                
                # 解析文档结构
                self.logger.info(f"[Parse] Begin parse -> {upload_path}")
                parsed_doc, created_md = await self._parse_document(upload_path, filename)
                parsed_documents.append(parsed_doc)
                # 如果产出了Markdown文件，纳入 files_to_create，供前端与会话保存
                if created_md and created_md.get("content"):
                    self.logger.info(f"[Parse] MarkItDown produced: {created_md.get('name')} (len={len(created_md.get('content',''))})")
                    files_to_create.append(created_md)
                # 汇总从Markdown中抽取的关键信息
                if parsed_doc.get("metadata", {}).get("extracted"):
                    aggregated_extracted[filename] = parsed_doc["metadata"]["extracted"]
                
                # 保存到wiki文件夹
                parsed_dir = "/root/project/git/project-agent/wiki"
                os.makedirs(parsed_dir, exist_ok=True)
                parsed_path = os.path.join(parsed_dir, f"{safe_filename}.json")
                with open(parsed_path, 'w', encoding='utf-8') as f:
                    json.dump(parsed_doc, f, ensure_ascii=False, indent=2)
            
            # 检查是否成功生成了wiki/招标文件.md
            wiki_tender_path = None
            for created_file in files_to_create:
                if created_file.get("name") == "招标文件.md" and created_file.get("type") == "wiki":
                    wiki_tender_path = "wiki/招标文件.md"
                    break
            
            response_content = f"成功解析了 {len(parsed_documents)} 个文档，解析结果已保存到wiki文件夹。"
            if wiki_tender_path:
                response_content += f"\n\n📄 **解析结果**: 已生成 `{wiki_tender_path}` 用于后续A-E工作流。"
            
            self.logger.info(f"[Summary] parsed_docs={len(parsed_documents)}, files_to_create={len(files_to_create)}")
            
            metadata = {
                "parsed_documents": parsed_documents,
                "parsed_files": [f"{doc['filename']}.json" for doc in parsed_documents],
                "files_to_create": files_to_create,
                # 直接输出两类核心要素，供后续阶段或UI使用
                "extracted_info": aggregated_extracted,
                "stage": "document_parsing",
                "action": "parsing_completed"
            }
            
            # 关键修复：如果成功生成了wiki/招标文件.md，更新tender_path用于后续工作流
            if wiki_tender_path:
                metadata["tender_path"] = wiki_tender_path
                metadata["next_stage"] = "bid_build_ready"
            
            return AgentResponse(
                content=response_content,
                metadata=metadata,
                status="completed"
            )
            
        except Exception as e:
            return AgentResponse(
                content=f"文档解析失败: {str(e)}",
                status="error"
            )
    
    async def _parse_document(self, file_path: str, filename: str) -> tuple[Dict[str, Any], Dict[str, Any]]:
        """解析单个文档，返回(结构化JSON, 生成的Markdown文件信息或None)"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._parse_pdf(file_path, filename)
        elif file_ext in ['.doc', '.docx']:
            return await self._parse_docx(file_path, filename)
        elif file_ext == '.txt':
            return await self._parse_txt(file_path, filename)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    async def _parse_pdf(self, file_path: str, filename: str) -> tuple[Dict[str, Any], Dict[str, Any]]:
        """解析PDF文档"""
        try:
            # 优先使用 microsoft/markitdown 将PDF转为Markdown
            self.logger.info(f"[MarkItDown] Converting PDF -> {file_path}")
            print(f"[MarkItDown] Converting PDF -> {file_path}")
            md_text = self._convert_to_markdown(file_path)
            created_md: Dict[str, Any] = {}
            if md_text:
                parsed_dir = "/root/project/git/project-agent/wiki"
                os.makedirs(parsed_dir, exist_ok=True)
                # 统一命名为"招标文件.md"
                base_name = "招标文件.md"
                md_path = os.path.join(parsed_dir, base_name)
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(md_text)
                self.logger.info(f"[MarkItDown] PDF->MD success: {md_path} (chars={len(md_text)})")
                print(f"[MarkItDown] PDF->MD success: {md_path} (chars={len(md_text)})")
                created_md = {"name": base_name, "content": md_text, "type": "wiki", "folder": "parsed"}

                # 使用LLM基于Markdown抽取结构与关键要素
                structure, extracted = await self._analyze_from_markdown(md_text, filename)
                return ({
                    "filename": filename,
                    "document_type": "PDF招标文件",
                    "structure": structure,
                    "metadata": {
                        "file_size": os.path.getsize(file_path),
                        "markdown_saved": True,
                        "markdown_file": base_name,
                        "extracted": extracted
                    }
                }, created_md)

            # 回退：直接提取文本
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content += f"\n--- 第 {page_num + 1} 页 ---\n"
                    text_content += page.extract_text()
                structure = await self._analyze_document_structure(text_content, filename)
                return ({
                    "filename": filename,
                    "document_type": "PDF招标文件",
                    "structure": structure,
                    "metadata": {
                        "total_pages": len(pdf_reader.pages),
                        "file_size": os.path.getsize(file_path),
                        "raw_text": text_content[:5000]
                    }
                }, {})
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")
    
    async def _parse_docx(self, file_path: str, filename: str) -> tuple[Dict[str, Any], Dict[str, Any]]:
        """解析DOCX文档"""
        try:
            # 使用 microsoft/markitdown 将Word转为Markdown
            self.logger.info(f"[MarkItDown] Converting DOC/DOCX -> {file_path}")
            print(f"[MarkItDown] Converting DOC/DOCX -> {file_path}")
            md_text = self._convert_to_markdown(file_path)
            created_md: Dict[str, Any] = {}
            if md_text:
                parsed_dir = "/root/project/git/project-agent/wiki"
                os.makedirs(parsed_dir, exist_ok=True)
                # 按需求统一命名为"招标文件.md"
                base_name = "招标文件.md"
                md_path = os.path.join(parsed_dir, base_name)
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(md_text)
                self.logger.info(f"[MarkItDown] DOCX->MD success: {md_path} (chars={len(md_text)})")
                print(f"[MarkItDown] DOCX->MD success: {md_path} (chars={len(md_text)})")
                created_md = {"name": base_name, "content": md_text, "type": "wiki", "folder": "parsed"}

                # 使用LLM基于Markdown抽取结构与关键要素
                structure, extracted = await self._analyze_from_markdown(md_text, filename)
                return ({
                    "filename": filename,
                    "document_type": "DOCX招标文件",
                    "structure": structure,
                    "metadata": {
                        "file_size": os.path.getsize(file_path),
                        "markdown_saved": True,
                        "markdown_file": base_name,
                        "extracted": extracted
                    }
                }, created_md)

            # 回退：python-docx 提取纯文本
            try:
                import docx
                doc = docx.Document(file_path)
                text_content = "\n".join(p.text for p in doc.paragraphs)
                structure = await self._analyze_document_structure(text_content, filename)
                return ({
                    "filename": filename,
                    "document_type": "DOCX招标文件",
                    "structure": structure,
                    "metadata": {
                        "total_paragraphs": len(doc.paragraphs),
                        "file_size": os.path.getsize(file_path),
                        "raw_text": text_content[:5000]
                    }
                }, {})
            except ImportError:
                self.logger.warning("python-docx not available, creating basic structure")
                return ({
                    "filename": filename,
                    "document_type": "DOCX招标文件",
                    "structure": [{
                        "level": 1,
                        "title": "文档内容",
                        "content": "需要安装python-docx来解析DOCX文件",
                        "subsections": []
                    }],
                    "metadata": {
                        "file_size": os.path.getsize(file_path),
                        "parse_method": "fallback"
                    }
                }, {})
        except Exception as e:
            raise Exception(f"DOCX解析失败: {str(e)}")
    
    async def _parse_txt(self, file_path: str, filename: str) -> tuple[Dict[str, Any], Dict[str, Any]]:
        """解析TXT文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # 简单文本分析（不使用LLM）
            lines = text_content.split('\n')
            sections = []
            current_section = None
            
            for line in lines:
                line = line.strip()
                if line.startswith('## '):
                    # 二级标题
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        "level": 2,
                        "title": line[3:],
                        "content": "",
                        "page": "N/A",
                        "subsections": []
                    }
                elif line.startswith('# '):
                    # 一级标题
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        "level": 1,
                        "title": line[2:],
                        "content": "",
                        "page": "N/A",
                        "subsections": []
                    }
                elif current_section and line:
                    # 添加到当前章节内容
                    current_section["content"] += line + "\n"
            
            if current_section:
                sections.append(current_section)
            
            # 如果没有找到章节，创建默认结构
            if not sections:
                sections = [{
                    "level": 1,
                    "title": "文档内容",
                    "content": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                    "page": "N/A",
                    "subsections": []
                }]
            
            return ({
                "filename": filename,
                "document_type": "TXT招标文件",
                "structure": sections,
                "metadata": {
                    "file_size": os.path.getsize(file_path),
                    "raw_text": text_content[:1000] + "..." if len(text_content) > 1000 else text_content
                }
            }, {})
        except Exception as e:
            raise Exception(f"TXT解析失败: {str(e)}")

    def _convert_to_markdown(self, file_path: str) -> str:
        """使用 microsoft/markitdown 将文件转换为Markdown。若不可用则返回空字符串。"""
        try:
            from markitdown import MarkItDown  # type: ignore
            md = MarkItDown()
            self.logger.info(f"[MarkItDown] convert() call for: {file_path}")
            print(f"[MarkItDown] convert() call for: {file_path}")
            result = md.convert(file_path)
            # 兼容不同版本的返回属性
            if hasattr(result, 'text_content'):
                return getattr(result, 'text_content') or ''
            if hasattr(result, 'markdown'):  # hypothetical
                return getattr(result, 'markdown') or ''
            if isinstance(result, str):
                return result
            # 尝试常见属性名
            for key in ('content', 'text'):
                val = getattr(result, key, '')
                if isinstance(val, str) and val:
                    return val
        except Exception as e:
            # 静默回退，以确保主流程不中断
            self.logger.warning(f"[MarkItDown] convert failed for {file_path}: {e}")
            print(f"[MarkItDown] convert failed for {file_path}: {e}")
        return ""
    
    async def _analyze_from_markdown(self, markdown_text: str, filename: str):
        """基于完整的 Markdown 文档：
        - 抽取文档结构
        - 提取两类关键信息：投标文件格式结构、技术规格书
        返回: (structure: List[dict], extracted: dict)
        """
        json_schema = """
{
  "structure": [...],
  "extracted": {
    "bid_format": {"found": true/false, "title": "", "content": ""},
    "tech_specifications": {"found": true/false, "title": "", "content": ""}
  }
}
"""

        analysis_prompt = (
            f"你将获得完整的招标文件（已转为Markdown）。\n"
            f"文件名：{filename}\n\n"
            "任务：\n"
            "1) 解析并返回文档的层级结构（structure，数组）。\n"
            "   - 每个节点包含: level(1/2/3..), title, content摘要(<=300字), subsections[]\n"
            "2) 提取两类关键信息（extracted 对象）：\n"
            "   - bid_format: 从"投标文件格式/投标文件内容/响应文件格式"等章节提取目录结构、装订顺序、章节要求。\n"
            "   - tech_specifications: 从"技术规格书/技术要求"等章节提取关键技术指标与约束。\n\n"
            "请严格输出一个JSON对象：\n"
            f"{json_schema}\n"
            "全文Markdown（截断展示，不要回显原文）：\n"
            f"{markdown_text[:8000]}\n"
        )
        
        try:
            result_text = await self.llm_client.generate([
                {"role": "system", "content": "你是文档结构与规则抽取专家。"},
                {"role": "user", "content": analysis_prompt}
            ])
            
            # 尝试解析JSON
            try:
                data = json.loads(result_text)
                structure = data.get('structure', []) if isinstance(data, dict) else []
                extracted = data.get('extracted', {}) if isinstance(data, dict) else {}
                if not isinstance(structure, list):
                    structure = []
                if not isinstance(extracted, dict):
                    extracted = {}
                return structure, extracted
            except:
                pass
            
            # 失败回退：仅返回单节点结构，extracted 为空
            return ([{
                "level": 1,
                "title": "文档内容",
                "content": "结构抽取失败，已保留原文用于后续处理",
                "subsections": []
            }], {})
            
        except Exception as e:
            # 返回基础结构
            return ([{
                "level": 1,
                "title": f"解析失败 - {filename}",
                "content": f"结构分析失败: {str(e)}",
                "subsections": []
            }], {})

    async def _analyze_document_structure(self, text_content: str, filename: str):
        """使用LLM分析文档结构"""
        try:
            analysis_prompt = f"""
请分析以下文档内容并识别其结构：

文件名：{filename}
文档内容：
{text_content[:3000]}...

请以JSON格式返回文档结构，包含章节标题、层级和内容摘要。
"""
            
            result_text = await self.llm_client.generate([
                {"role": "system", "content": "你是文档结构分析专家。"},
                {"role": "user", "content": analysis_prompt}
            ])
            
            # 尝试解析JSON，如果失败则返回基础结构
            try:
                return json.loads(result_text)
            except:
                return [{
                    "level": 1,
                    "title": "文档内容",
                    "content": text_content[:300] + "..." if len(text_content) > 300 else text_content,
                    "subsections": []
                }]
        except Exception as e:
            return [{
                "level": 1,
                "title": f"解析失败 - {filename}",
                "content": f"结构分析失败: {str(e)}",
                "subsections": []
            }]